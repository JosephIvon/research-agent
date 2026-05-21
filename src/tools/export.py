"""多格式导出工具"""


def export_to_pdf(markdown_content: str, output_path: str) -> bool:
    """将 Markdown 导出为 PDF"""
    try:
        import markdown
        from weasyprint import HTML, CSS
    except ImportError as e:
        raise ImportError(f"PDF export requires weasyprint: {e}")

    # Convert markdown to HTML
    html_content = markdown.markdown(
        markdown_content,
        extensions=['tables', 'fenced_code']
    )

    # Wrap in full HTML document with styling
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: 'SimSun', 'Microsoft YaHei', sans-serif;
            padding: 40px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 8px; }}
        h2 {{ color: #34495e; margin-top: 1.5em; border-bottom: 1px solid #eee; padding-bottom: 4px; }}
        h3 {{ color: #555; margin-top: 1.2em; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #fafafa; }}
        code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-family: monospace; }}
        pre {{ background: #f4f4f4; padding: 1em; border-radius: 6px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 1em 0; padding-left: 1em; color: #666; }}
        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
        li {{ margin: 0.3em 0; }}
    </style>
</head>
<body>{html_content}</body>
</html>"""

    HTML(string=full_html).write_pdf(output_path)
    return True


def export_to_word(markdown_content: str, output_path: str) -> bool:
    """将 Markdown 导出为 Word (.docx)"""
    try:
        from docx import Document
        import markdown
    except ImportError as e:
        raise ImportError(f"Word export requires python-docx: {e}")

    doc = Document()

    # Parse markdown line by line (basic conversion)
    lines = markdown_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph('')
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.add_run(line[2:]).italic = True
        elif line.startswith('|'):
            # Skip table lines for basic implementation
            continue
        else:
            # Handle inline code
            if line.startswith('`') and line.endswith('`'):
                p = doc.add_paragraph()
                p.add_run(line[1:-1]).font.name = 'Courier New'
            else:
                doc.add_paragraph(line)

    doc.save(output_path)
    return True


def export_to_html(markdown_content: str, output_path: str) -> bool:
    """将 Markdown 导出为 HTML"""
    import markdown

    html_content = markdown.markdown(
        markdown_content,
        extensions=['tables', 'fenced_code', 'toc']
    )

    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Research Report</title>
    <style>
        body {{ font-family: system-ui, -apple-system, sans-serif; max-width: 900px; margin: 0 auto; padding: 2rem; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 8px; }}
        h2 {{ color: #34495e; margin-top: 1.5em; }}
        h3 {{ color: #555; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; }}
        th {{ background: #f2f2f2; font-weight: bold; }}
        pre {{ background: #f4f4f4; padding: 1rem; border-radius: 6px; overflow-x: auto; }}
        code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 1em 0; padding-left: 1em; color: #666; }}
        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
    </style>
</head>
<body>{html_content}</body>
</html>"""

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    return True
